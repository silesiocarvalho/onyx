#!/usr/bin/env python3
"""
report_generator.py — Multi-format Report Generator
Produces CSV, Excel (multi-sheet), and PDF from an enriched audit JSON
(output of ai_analyzer.py or raw audit_tool.py output).
"""

import csv
import io
import json
import os
import textwrap
from datetime import datetime
from typing import Optional

# ---------------------------------------------------------------------------
# Optional imports
# ---------------------------------------------------------------------------
try:
    import openpyxl
    from openpyxl.styles import (Alignment, Border, Font, GradientFill,
                                  PatternFill, Side)
    from openpyxl.chart import BarChart, PieChart, Reference
    from openpyxl.chart.series import DataPoint
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm, inch, mm
    from reportlab.platypus import (BaseDocTemplate, Frame, HRFlowable,
                                     KeepTogether, PageBreak, PageTemplate,
                                     Paragraph, SimpleDocTemplate, Spacer,
                                     Table, TableStyle)
    from reportlab.platypus.tableofcontents import TableOfContents
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
COLORS = {
    "pass":    "#2ECC71",
    "fail":    "#E74C3C",
    "manual":  "#F39C12",
    "skipped": "#95A5A6",
    "error":   "#C0392B",
    "critical":"#8E44AD",
    "high":    "#E74C3C",
    "medium":  "#F39C12",
    "low":     "#27AE60",
    "info":    "#3498DB",
    "header":  "#1A252F",
    "accent":  "#2980B9",
    "white":   "#FFFFFF",
    "light":   "#ECF0F1",
    "mid":     "#BDC3C7",
}

STATUS_COLOR = {
    "PASS":           COLORS["pass"],
    "FAIL":           COLORS["fail"],
    "MANUAL":         COLORS["manual"],
    "SKIPPED":        COLORS["skipped"],
    "ERROR":          COLORS["error"],
    "RECOMMENDATION": COLORS["info"],
}

RISK_COLOR = {
    "Critical": COLORS["critical"],
    "High":     COLORS["high"],
    "Medium":   COLORS["medium"],
    "Low":      COLORS["low"],
    "Info":     COLORS["info"],
    "Unknown":  COLORS["mid"],
}

def hex_to_rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


class _CISDocTemplate(BaseDocTemplate):
    """BaseDocTemplate subclass that feeds section headings into the TOC."""

    def __init__(self, filename, **kwargs):
        super().__init__(filename, **kwargs)
        top_m    = kwargs.get("topMargin",    1.8 * cm)
        bot_m    = kwargs.get("bottomMargin", 1.4 * cm)
        left_m   = kwargs.get("leftMargin",   1.6 * cm)
        right_m  = kwargs.get("rightMargin",  1.6 * cm)
        w = self.pagesize[0] - left_m - right_m
        h = self.pagesize[1] - top_m  - bot_m
        frame = Frame(left_m, bot_m, w, h, id="main", showBoundary=0)
        self.addPageTemplates([PageTemplate(id="main", frames=[frame],
                                            onPage=self._hf)])
        self._hf_cb = None

    def set_header_footer(self, cb):
        self._hf_cb = cb

    def _hf(self, canvas, doc):
        if self._hf_cb:
            self._hf_cb(canvas, doc)

    def afterFlowable(self, flowable):
        if not isinstance(flowable, Paragraph):
            return
        sname = getattr(flowable.style, "name", "")
        if sname == "section_header":
            self.notify("TOCEntry", (0, flowable.getPlainText(), self.page))
        elif sname == "heading2":
            self.notify("TOCEntry", (1, flowable.getPlainText(), self.page))


# ---------------------------------------------------------------------------
# Shared data preparation
# ---------------------------------------------------------------------------
def _prepare(data: dict) -> dict:
    """Normalise raw or enriched JSON into a unified working dict."""
    from tools.compliance_mappings import annotate_findings

    findings  = data.get("results", [])
    meta      = data.get("meta", {})
    stats     = data.get("summary", {})
    narrative = data.get("narrative", {})
    chains    = data.get("attack_chains", [])
    frameworks = meta.get("frameworks", [])

    # Annotate findings with compliance mappings if frameworks selected
    if frameworks:
        findings = annotate_findings(findings, frameworks)

    # Exclude RECOMMENDATION from score/stats — they are advisory, not audit checks
    audit_findings = [f for f in findings if f.get("status") != "RECOMMENDATION"]
    total   = len(audit_findings)
    passes  = sum(1 for f in audit_findings if f["status"] == "PASS")
    fails   = sum(1 for f in audit_findings if f["status"] == "FAIL")
    manual  = sum(1 for f in audit_findings if f["status"] == "MANUAL")
    skipped = sum(1 for f in audit_findings if f["status"] == "SKIPPED")
    errors  = sum(1 for f in audit_findings if f["status"] == "ERROR")
    score   = round(passes / total * 100, 1) if total else 0.0

    return dict(
        findings=findings, meta=meta, stats=stats,
        narrative=narrative, chains=chains, frameworks=frameworks,
        total=total, passes=passes, fails=fails,
        manual=manual, skipped=skipped, errors=errors, score=score,
    )


# ============================================================================
# 1. CSV REPORT
# ============================================================================
def generate_csv(data: dict, output_path: str) -> str:
    p = _prepare(data)

    fieldnames = [
        "control_id", "description", "level", "status",
        "expected", "actual", "remediation", "notes",
        "ai_risk_level", "ai_business_impact", "ai_attack_scenario",
        "ai_remediation_effort", "ai_priority_rank", "ai_cve_reference",
        "timestamp",
        "nist_csf", "iso_27001", "pci_dss", "mitre_attack",
    ]

    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for finding in p["findings"]:
            if finding.get("status") == "RECOMMENDATION":
                continue
            ai         = finding.get("ai_analysis", {})
            compliance = finding.get("compliance", {})
            row = {
                "control_id":            finding.get("control_id", ""),
                "description":           finding.get("description", ""),
                "level":                 finding.get("level", ""),
                "status":                finding.get("status", ""),
                "expected":              finding.get("expected", ""),
                "actual":                str(finding.get("actual", ""))[:200],
                "remediation":           finding.get("remediation", ""),
                "notes":                 finding.get("notes", ""),
                "ai_risk_level":         ai.get("risk_level", ""),
                "ai_business_impact":    ai.get("business_impact", ""),
                "ai_attack_scenario":    ai.get("attack_scenario", ""),
                "ai_remediation_effort": ai.get("remediation_effort", ""),
                "ai_priority_rank":      ai.get("priority_rank", ""),
                "ai_cve_reference":      ai.get("cve_or_reference", ""),
                "timestamp":             finding.get("timestamp", ""),
                "nist_csf":              " | ".join(compliance.get("nist_csf", [])),
                "iso_27001":             " | ".join(compliance.get("iso_27001", [])),
                "pci_dss":               " | ".join(compliance.get("pci_dss", [])),
                "mitre_attack":          " | ".join(compliance.get("mitre_attack", [])),
            }
            writer.writerow(row)

    print(f"[REPORT] ✓ CSV  → {os.path.abspath(output_path)}")
    return output_path


# ============================================================================
# 2. EXCEL REPORT
# ============================================================================
def _xl_hex(hex_color: str) -> str:
    """openpyxl needs ARGB without #."""
    return "FF" + hex_color.lstrip("#").upper()

def _xl_fill(hex_color: str) -> "PatternFill":
    return PatternFill("solid", fgColor=_xl_hex(hex_color))

def _xl_font(bold=False, color="#000000", size=11, italic=False) -> "Font":
    return Font(bold=bold, color=_xl_hex(color), size=size, italic=italic)

def _xl_border() -> "Border":
    side = Side(style="thin", color="FFBDC3C7")
    return Border(left=side, right=side, top=side, bottom=side)

def _xl_align(h="left", v="center", wrap=False) -> "Alignment":
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)


def _sheet_dashboard(wb, p):
    ws = wb.create_sheet("📊 Dashboard", 0)
    ws.sheet_view.showGridLines = False

    # Column widths
    for col, w in zip("ABCDEFG", [2, 20, 14, 14, 14, 14, 20]):
        ws.column_dimensions[get_column_letter(ord(col) - ord('A') + 1)].width = w

    # Title block
    ws.row_dimensions[1].height = 8
    ws.row_dimensions[2].height = 40
    ws.row_dimensions[3].height = 24
    ws.merge_cells("B2:G2")
    ws["B2"] = "FW AI Audit Security Assessment"
    ws["B2"].font      = _xl_font(bold=True, size=18, color=COLORS["white"])
    ws["B2"].fill      = _xl_fill(COLORS["header"])
    ws["B2"].alignment = _xl_align("center")

    ws.merge_cells("B3:G3")
    target = p["meta"].get("target", "Unknown")
    ts     = p["meta"].get("generated", "")[:19].replace("T", " ")
    benchmark = p["meta"].get("benchmark", "FW AI Audit Security Assessment")
    ws["B3"] = f"Target: {target}   |   Assessment: {ts}   |   {benchmark}"
    ws["B3"].font      = _xl_font(size=10, color=COLORS["white"])
    ws["B3"].fill      = _xl_fill(COLORS["accent"])
    ws["B3"].alignment = _xl_align("center")

    # Score block
    ws.row_dimensions[5].height = 60
    ws.merge_cells("B5:C5")
    ws["B5"] = f"{p['score']}%"
    ws["B5"].font      = _xl_font(bold=True, size=36, color=COLORS["header"])
    ws["B5"].alignment = _xl_align("center")

    ws.merge_cells("D5:E5")
    overall = p["narrative"].get("overall_risk_rating", "Unknown")
    risk_c  = RISK_COLOR.get(overall, COLORS["mid"])
    ws["D5"] = f"Overall Risk: {overall}"
    ws["D5"].font      = _xl_font(bold=True, size=14, color=COLORS["white"])
    ws["D5"].fill      = _xl_fill(risk_c)
    ws["D5"].alignment = _xl_align("center")

    ws.merge_cells("B6:E6")
    ws["B6"] = "Compliance Score"
    ws["B6"].font      = _xl_font(bold=True, size=12, color=COLORS["header"])
    ws["B6"].alignment = _xl_align("center")

    # Stat boxes  row 8-10
    ws.row_dimensions[8].height = 14
    stats_data = [
        ("PASS",    p["passes"],  COLORS["pass"]),
        ("FAIL",    p["fails"],   COLORS["fail"]),
        ("MANUAL",  p["manual"],  COLORS["manual"]),
        ("SKIPPED", p["skipped"], COLORS["skipped"]),
    ]
    cols = ["B", "C", "D", "E"]
    ws.row_dimensions[9].height = 36
    ws.row_dimensions[10].height = 20
    for col_letter, (label, count, color) in zip(cols, stats_data):
        col_idx = ord(col_letter) - ord('A') + 1
        # Value
        ws.cell(row=9, column=col_idx).value     = count
        ws.cell(row=9, column=col_idx).font      = _xl_font(bold=True, size=22, color=COLORS["white"])
        ws.cell(row=9, column=col_idx).fill      = _xl_fill(color)
        ws.cell(row=9, column=col_idx).alignment = _xl_align("center")
        # Label
        ws.cell(row=10, column=col_idx).value     = label
        ws.cell(row=10, column=col_idx).font      = _xl_font(bold=True, size=10, color=COLORS["white"])
        ws.cell(row=10, column=col_idx).fill      = _xl_fill(color)
        ws.cell(row=10, column=col_idx).alignment = _xl_align("center")

    # Pie chart data (hidden, used by chart)
    ws.row_dimensions[14].height = 14
    ws["B14"] = "Status"
    ws["C14"] = "Count"
    ws["B14"].font = _xl_font(bold=True)
    ws["C14"].font = _xl_font(bold=True)
    chart_rows = [("Pass", p["passes"]), ("Fail", p["fails"]),
                  ("Manual", p["manual"]), ("Skipped", p["skipped"])]
    for i, (label, count) in enumerate(chart_rows, start=15):
        ws.cell(row=i, column=2).value = label
        ws.cell(row=i, column=3).value = count

    pie = PieChart()
    pie.title = "Compliance Breakdown"
    pie.style = 10
    pie.width  = 14
    pie.height = 10
    labels = Reference(ws, min_col=2, min_row=15, max_row=18)
    data   = Reference(ws, min_col=3, min_row=14, max_row=18)
    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    slice_colors = ["00B050", "FF0000", "FFC000", "808080"]
    for idx, color in enumerate(slice_colors):
        pt = DataPoint(idx=idx)
        pt.graphicalProperties.solidFill = color
        pie.series[0].dPt.append(pt)
    ws.add_chart(pie, "F5")

    # Executive summary block
    exec_sum = p["narrative"].get("executive_summary", {})
    headline  = exec_sum.get("headline", "")
    para1     = exec_sum.get("paragraph_1", "")

    row = 21
    ws.row_dimensions[row].height = 18
    ws.merge_cells(f"B{row}:G{row}")
    ws[f"B{row}"] = "EXECUTIVE SUMMARY"
    ws[f"B{row}"].font      = _xl_font(bold=True, size=12, color=COLORS["white"])
    ws[f"B{row}"].fill      = _xl_fill(COLORS["header"])
    ws[f"B{row}"].alignment = _xl_align("center")

    row += 1
    ws.row_dimensions[row].height = 18
    ws.merge_cells(f"B{row}:G{row}")
    ws[f"B{row}"] = headline
    ws[f"B{row}"].font      = _xl_font(bold=True, italic=True, size=11)
    ws[f"B{row}"].alignment = _xl_align(wrap=True)

    row += 1
    ws.row_dimensions[row].height = 60
    ws.merge_cells(f"B{row}:G{row}")
    ws[f"B{row}"] = para1
    ws[f"B{row}"].font      = _xl_font(size=10)
    ws[f"B{row}"].alignment = _xl_align(wrap=True)

    # Top 5 actions
    row += 2
    ws.row_dimensions[row].height = 18
    ws.merge_cells(f"B{row}:G{row}")
    ws[f"B{row}"] = "TOP PRIORITY ACTIONS"
    ws[f"B{row}"].font      = _xl_font(bold=True, size=12, color=COLORS["white"])
    ws[f"B{row}"].fill      = _xl_fill(COLORS["header"])
    ws[f"B{row}"].alignment = _xl_align("center")

    actions = p["narrative"].get("top_5_priority_actions", [])
    headers = ["#", "Action", "Justification", "Effort", "Impact"]
    row += 1
    for ci, h in enumerate(headers, start=2):
        ws.row_dimensions[row].height = 16
        ws.cell(row=row, column=ci).value     = h
        ws.cell(row=row, column=ci).font      = _xl_font(bold=True, color=COLORS["white"])
        ws.cell(row=row, column=ci).fill      = _xl_fill(COLORS["accent"])
        ws.cell(row=row, column=ci).alignment = _xl_align("center")

    for action in actions:
        row += 1
        effort = action.get("effort", "")
        effort_color = {"Low": COLORS["pass"], "Medium": COLORS["manual"],
                        "High": COLORS["fail"]}.get(effort, COLORS["mid"])
        vals = [action.get("rank", ""), action.get("action", ""),
                action.get("justification", ""), effort, action.get("impact", "")]
        for ci, val in enumerate(vals, start=2):
            ws.row_dimensions[row].height = 36
            cell = ws.cell(row=row, column=ci)
            cell.value     = val
            cell.alignment = _xl_align(wrap=True)
            cell.border    = _xl_border()
            if ci == 5 and effort_color:
                cell.fill = _xl_fill(effort_color)
                cell.font = _xl_font(color=COLORS["white"], bold=True)


def _sheet_findings(wb, p):
    ws = wb.create_sheet("🔍 All Findings")
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A2"

    col_widths = [8, 45, 8, 12, 14, 22, 22, 35, 12, 16]
    col_letters = [get_column_letter(i) for i in range(1, len(col_widths) + 1)]
    for letter, width in zip(col_letters, col_widths):
        ws.column_dimensions[letter].width = width

    headers = ["ID", "Description", "Level", "Status",
               "Risk Level", "Expected", "Actual", "Remediation",
               "Effort", "Priority"]
    ws.row_dimensions[1].height = 20
    for ci, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=ci)
        cell.value     = h
        cell.font      = _xl_font(bold=True, size=11, color=COLORS["white"])
        cell.fill      = _xl_fill(COLORS["header"])
        cell.alignment = _xl_align("center")
        cell.border    = _xl_border()

    for row_idx, finding in enumerate(p["findings"], start=2):
        ai      = finding.get("ai_analysis", {})
        status  = finding.get("status", "")
        risk    = ai.get("risk_level", "")
        s_color = STATUS_COLOR.get(status, COLORS["mid"])
        r_color = RISK_COLOR.get(risk, COLORS["mid"])
        bg      = COLORS["light"] if row_idx % 2 == 0 else COLORS["white"]

        row_data = [
            finding.get("control_id", ""),
            finding.get("description", ""),
            finding.get("level", ""),
            status,
            risk,
            str(finding.get("expected", "") or "")[:60],
            str(finding.get("actual", "") or "")[:60],
            finding.get("remediation", ""),
            ai.get("remediation_effort", ""),
            ai.get("priority_rank", ""),
        ]

        ws.row_dimensions[row_idx].height = 30
        for ci, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=ci)
            cell.value     = val
            cell.alignment = _xl_align(wrap=True)
            cell.border    = _xl_border()
            cell.fill      = _xl_fill(bg)
            # Status column
            if ci == 4:
                cell.fill = _xl_fill(s_color)
                cell.font = _xl_font(bold=True, color=COLORS["white"])
                cell.alignment = _xl_align("center")
            # Risk column
            elif ci == 5 and risk:
                cell.fill = _xl_fill(r_color)
                cell.font = _xl_font(bold=True, color=COLORS["white"])
                cell.alignment = _xl_align("center")
            elif ci == 9:
                effort_c = {"Low": COLORS["pass"], "Medium": COLORS["manual"],
                            "High": COLORS["fail"]}.get(str(val), bg)
                cell.fill = _xl_fill(effort_c)
                if str(val) in ("Low", "Medium", "High"):
                    cell.font = _xl_font(color=COLORS["white"])
                cell.alignment = _xl_align("center")

    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"


def _sheet_attack_chains(wb, p):
    if not p["chains"]:
        return
    ws = wb.create_sheet("⛓️ Attack Chains")
    ws.sheet_view.showGridLines = False

    col_ws = [8, 28, 12, 24, 55, 45, 30]
    for i, w in enumerate(col_ws, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    headers = ["ID", "Chain Name", "Risk", "Controls Involved",
               "Attack Narrative", "Blast Radius", "Fix Order"]
    ws.row_dimensions[1].height = 20
    for ci, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=ci)
        cell.value     = h
        cell.font      = _xl_font(bold=True, color=COLORS["white"])
        cell.fill      = _xl_fill(COLORS["header"])
        cell.alignment = _xl_align("center")
        cell.border    = _xl_border()

    for row_idx, chain in enumerate(p["chains"], start=2):
        risk    = chain.get("risk_level", "")
        r_color = RISK_COLOR.get(risk, COLORS["mid"])
        bg      = COLORS["light"] if row_idx % 2 == 0 else COLORS["white"]

        fix_order = chain.get("priority_fix_order", [])
        if isinstance(fix_order, list):
            fix_order = " → ".join(fix_order)
        controls = chain.get("controls_involved", [])
        if isinstance(controls, list):
            controls = ", ".join(controls)

        row_data = [
            chain.get("chain_id", ""),
            chain.get("chain_name", ""),
            risk,
            controls,
            chain.get("attack_narrative", ""),
            chain.get("blast_radius", ""),
            fix_order,
        ]
        ws.row_dimensions[row_idx].height = 60
        for ci, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=ci)
            cell.value     = val
            cell.alignment = _xl_align(wrap=True)
            cell.border    = _xl_border()
            cell.fill      = _xl_fill(bg)
            if ci == 3:
                cell.fill = _xl_fill(r_color)
                cell.font = _xl_font(bold=True, color=COLORS["white"])
                cell.alignment = _xl_align("center")


def _sheet_remediation(wb, p):
    ws = wb.create_sheet("🛠️ Remediation Roadmap")
    ws.sheet_view.showGridLines = False

    col_ws = [6, 8, 35, 10, 10, 12, 45, 50]
    for i, w in enumerate(col_ws, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    headers = ["Priority", "ID", "Description", "Level",
               "Status", "Effort", "Remediation Steps", "Business Impact"]
    ws.row_dimensions[1].height = 20
    for ci, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=ci)
        cell.value     = h
        cell.font      = _xl_font(bold=True, color=COLORS["white"])
        cell.fill      = _xl_fill(COLORS["accent"])
        cell.alignment = _xl_align("center")
        cell.border    = _xl_border()

    # Sort by priority_rank
    actionable = [f for f in p["findings"] if f["status"] in ("FAIL", "MANUAL")]
    def sort_key(f):
        try:
            return int(f.get("ai_analysis", {}).get("priority_rank", 999))
        except (TypeError, ValueError):
            return 999
    actionable.sort(key=sort_key)

    for row_idx, finding in enumerate(actionable, start=2):
        ai      = finding.get("ai_analysis", {})
        status  = finding.get("status", "")
        effort  = ai.get("remediation_effort", "")
        steps   = ai.get("remediation_steps", finding.get("remediation", ""))
        if isinstance(steps, list):
            steps = "\n".join(f"{i+1}. {s}" for i, s in enumerate(steps))

        s_color = STATUS_COLOR.get(status, COLORS["mid"])
        e_color = {"Low": COLORS["pass"], "Medium": COLORS["manual"],
                   "High": COLORS["fail"]}.get(effort, COLORS["mid"])
        bg = COLORS["light"] if row_idx % 2 == 0 else COLORS["white"]

        row_data = [
            ai.get("priority_rank", ""),
            finding.get("control_id", ""),
            finding.get("description", ""),
            finding.get("level", ""),
            status,
            effort,
            steps,
            ai.get("business_impact", ""),
        ]
        ws.row_dimensions[row_idx].height = 50
        for ci, val in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=ci)
            cell.value     = val
            cell.alignment = _xl_align(wrap=True)
            cell.border    = _xl_border()
            cell.fill      = _xl_fill(bg)
            if ci == 5:
                cell.fill = _xl_fill(s_color)
                cell.font = _xl_font(bold=True, color=COLORS["white"])
                cell.alignment = _xl_align("center")
            elif ci == 6:
                cell.fill = _xl_fill(e_color)
                if effort:
                    cell.font = _xl_font(bold=True, color=COLORS["white"])
                cell.alignment = _xl_align("center")


def _sheet_raw(wb, p):
    ws = wb.create_sheet("📄 Raw Data")
    headers = ["control_id", "description", "level", "status",
               "expected", "actual", "remediation", "notes",
               "ai_risk_level", "ai_business_impact", "ai_attack_scenario",
               "ai_remediation_effort", "ai_priority_rank", "timestamp"]
    ws.append(headers)
    for f in p["findings"]:
        if f.get("status") == "RECOMMENDATION":
            continue
        ai = f.get("ai_analysis", {})
        ws.append([
            f.get("control_id", ""), f.get("description", ""),
            f.get("level", ""),      f.get("status", ""),
            str(f.get("expected", "") or "")[:100],
            str(f.get("actual", "") or "")[:100],
            f.get("remediation", ""),  f.get("notes", ""),
            ai.get("risk_level", ""),  ai.get("business_impact", ""),
            ai.get("attack_scenario", ""), ai.get("remediation_effort", ""),
            ai.get("priority_rank", ""), f.get("timestamp", ""),
        ])


def _sheet_compliance(wb, p):
    """Compliance mapping sheet — only added when frameworks were selected."""
    from tools.compliance_mappings import FRAMEWORKS, NIST_CSF_LABELS, MITRE_ATTACK_LABELS

    frameworks = p.get("frameworks", [])
    if not frameworks:
        return

    ws = wb.create_sheet("🛡️ Compliance")
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A2"

    fw_names = [FRAMEWORKS[fw]["name"] for fw in frameworks if fw in FRAMEWORKS]

    # Dynamic column widths: ID + Description + one per framework
    col_widths = [8, 42] + [30] * len(frameworks)
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Header
    headers = ["ID", "Description"] + fw_names
    ws.row_dimensions[1].height = 22
    for ci, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=ci)
        cell.value     = h
        cell.font      = _xl_font(bold=True, size=11, color=COLORS["white"])
        cell.fill      = _xl_fill(COLORS["header"])
        cell.alignment = _xl_align("center")
        cell.border    = _xl_border()

    fw_colors = {
        "nist_csf":     "1A5276",
        "iso_27001":    "1E8449",
        "pci_dss":      "7D3C98",
        "mitre_attack": "922B21",
    }

    for row_idx, finding in enumerate(p["findings"], start=2):
        compliance = finding.get("compliance", {})
        bg = COLORS["light"] if row_idx % 2 == 0 else COLORS["white"]
        ws.row_dimensions[row_idx].height = 40

        # ID
        c = ws.cell(row=row_idx, column=1)
        c.value     = finding.get("control_id", "")
        c.font      = _xl_font(bold=True, size=10)
        c.alignment = _xl_align("center")
        c.fill      = _xl_fill(bg)
        c.border    = _xl_border()

        # Status colouring on ID cell
        status = finding.get("status", "")
        if status in STATUS_COLOR:
            c.fill = _xl_fill(STATUS_COLOR[status])
            c.font = _xl_font(bold=True, size=10, color=COLORS["white"])

        # Description
        d = ws.cell(row=row_idx, column=2)
        d.value     = finding.get("description", "")
        d.font      = _xl_font(size=10)
        d.alignment = _xl_align(wrap=True)
        d.fill      = _xl_fill(bg)
        d.border    = _xl_border()

        # One column per selected framework
        for fi, fw in enumerate(frameworks, start=3):
            controls = compliance.get(fw, [])
            cell = ws.cell(row=row_idx, column=fi)
            cell.value     = "\n".join(controls) if controls else "—"
            cell.font      = _xl_font(size=9, color=fw_colors.get(fw, "000000") if controls else "999999")
            cell.alignment = _xl_align(wrap=True)
            cell.fill      = _xl_fill(bg)
            cell.border    = _xl_border()

    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"


def generate_excel(data: dict, output_path: str) -> str:
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl required: pip install openpyxl --break-system-packages")

    p  = _prepare(data)
    wb = openpyxl.Workbook()
    # Remove default sheet
    wb.remove(wb.active)

    _sheet_dashboard(wb, p)
    _sheet_findings(wb, p)
    _sheet_attack_chains(wb, p)
    _sheet_remediation(wb, p)
    _sheet_compliance(wb, p)
    _sheet_raw(wb, p)

    wb.save(output_path)
    print(f"[REPORT] ✓ Excel → {os.path.abspath(output_path)}")
    return output_path


# ============================================================================
# 3. PDF REPORT
# ============================================================================
def _rl_color(hex_color: str):
    r, g, b = hex_to_rgb(hex_color)
    return colors.Color(r / 255, g / 255, b / 255)


def _strip_html(text: str) -> str:
    """Remove HTML tags from AI-generated text."""
    import re
    return re.sub(r'<[^>]+>', '', str(text))

def _esc(text: str) -> str:
    """Escape XML special chars in a variable to safely embed in ReportLab markup."""
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def _para(text: str, style) -> "Paragraph":
    """Paragraph from AI/user content — strips HTML tags, escapes remaining XML chars."""
    text = _strip_html(str(text)).replace("&", "&amp;")
    return Paragraph(text, style)


def _split_numbered(text: str) -> list:
    """Split AI-generated text into chunks on natural numbered boundaries.

    Handles three common AI output patterns:
      1. "1. sentence. 2. sentence." — numbered with period
      2. "(1) sentence (2) sentence" — numbered with parens
      3. Flowing prose — split on '. ' before a capital letter
    Returns a list of strings; at minimum [original_text].
    """
    import re
    text = _strip_html(str(text)).strip()
    if not text:
        return [text]

    # Pattern 1: ". N. " — sentence ends, next numbered item begins
    if re.search(r'(?<=[.!?])\s+\d+\.\s', text):
        parts = re.split(r'(?<=[.!?])\s+(?=\d+\.\s)', text)
        if len(parts) > 1:
            return [p.strip() for p in parts if p.strip()]

    # Pattern 2: "(N)" numbered items
    if re.search(r'\(\d+\)\s', text):
        parts = re.split(r'\s+(?=\(\d+\)\s)', text)
        if len(parts) > 1:
            return [p.strip() for p in parts if p.strip()]

    # Pattern 3: sentence boundaries before a capital letter
    parts = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    if len(parts) > 1:
        return [p.strip() for p in parts if p.strip()]

    return [text]


def _multi_para(text: str, style, gap: float = 0.12) -> list:
    """Return a list of Paragraph + Spacer flowables, one per natural paragraph in text."""
    chunks = _split_numbered(text)
    result = []
    for i, chunk in enumerate(chunks):
        result.append(Paragraph(chunk.replace("&", "&amp;"), style))
        if i < len(chunks) - 1:
            result.append(Spacer(1, gap * cm))
    return result


def _build_styles() -> dict:
    base = getSampleStyleSheet()
    styles = {}

    styles["title"] = ParagraphStyle("title",
        parent=base["Normal"],
        fontSize=22, fontName="Helvetica-Bold",
        leading=28,
        textColor=_rl_color(COLORS["white"]),
        alignment=TA_CENTER, spaceAfter=6)

    styles["subtitle"] = ParagraphStyle("subtitle",
        parent=base["Normal"],
        fontSize=12, fontName="Helvetica",
        textColor=_rl_color(COLORS["white"]),
        alignment=TA_CENTER, spaceAfter=4)

    styles["section_header"] = ParagraphStyle("section_header",
        parent=base["Normal"],
        fontSize=14, fontName="Helvetica-Bold",
        textColor=_rl_color(COLORS["white"]),
        backColor=_rl_color(COLORS["header"]),
        alignment=TA_LEFT, spaceBefore=20, spaceAfter=10,
        leftIndent=6, rightIndent=6, leading=22)

    styles["body"] = ParagraphStyle("body",
        parent=base["Normal"],
        fontSize=9.5, fontName="Helvetica",
        textColor=_rl_color(COLORS["header"]),
        alignment=TA_JUSTIFY, spaceAfter=10, leading=16)

    styles["body_bold"] = ParagraphStyle("body_bold",
        parent=styles["body"],
        fontName="Helvetica-Bold", spaceAfter=8)

    styles["body_italic"] = ParagraphStyle("body_italic",
        parent=styles["body"],
        fontName="Helvetica-Oblique", spaceAfter=8)

    styles["field_label"] = ParagraphStyle("field_label",
        parent=base["Normal"],
        fontSize=9, fontName="Helvetica-Bold",
        textColor=_rl_color(COLORS["accent"]),
        spaceBefore=8, spaceAfter=2)

    styles["field_value"] = ParagraphStyle("field_value",
        parent=base["Normal"],
        fontSize=9, fontName="Helvetica",
        textColor=_rl_color(COLORS["header"]),
        spaceAfter=6, leading=14, leftIndent=6)

    styles["label"] = ParagraphStyle("label",
        parent=base["Normal"],
        fontSize=8, fontName="Helvetica-Bold",
        textColor=_rl_color(COLORS["accent"]),
        spaceAfter=2)

    styles["small"] = ParagraphStyle("small",
        parent=base["Normal"],
        fontSize=8, fontName="Helvetica",
        textColor=_rl_color(COLORS["header"]),
        leading=13)

    styles["small_white"] = ParagraphStyle("small_white",
        parent=base["Normal"],
        fontSize=8, fontName="Helvetica-Bold",
        textColor=colors.white,
        leading=13)

    styles["heading2"] = ParagraphStyle("heading2",
        parent=base["Normal"],
        fontSize=12, fontName="Helvetica-Bold",
        textColor=_rl_color(COLORS["header"]),
        spaceBefore=16, spaceAfter=8)

    styles["caption"] = ParagraphStyle("caption",
        parent=base["Normal"],
        fontSize=8, fontName="Helvetica-Oblique",
        textColor=_rl_color(COLORS["mid"]),
        alignment=TA_CENTER)

    # TOC entry styles
    styles["toc_h1"] = ParagraphStyle("toc_h1",
        parent=base["Normal"],
        fontSize=10, fontName="Helvetica-Bold",
        textColor=_rl_color(COLORS["header"]),
        spaceBefore=4, spaceAfter=2, leftIndent=0)

    styles["toc_h2"] = ParagraphStyle("toc_h2",
        parent=base["Normal"],
        fontSize=9, fontName="Helvetica",
        textColor=_rl_color(COLORS["accent"]),
        spaceBefore=2, spaceAfter=1, leftIndent=12)

    return styles


def _stat_table(p: dict, styles: dict) -> "Table":
    """Compact compliance stat table for the cover/summary page."""
    data = [
        ["Metric", "Count", "Metric", "Count"],
        ["Pass",    p["passes"],  "Manual",  p["manual"]],
        ["Fail",    p["fails"],   "Skipped",  p["skipped"]],
        ["Score",   f"{p['score']}%", "Total Checks", p["total"]],
    ]
    t = Table(data, colWidths=[2.5*cm, 1.5*cm, 2.5*cm, 1.5*cm])
    ts_style = TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0),  _rl_color(COLORS["header"])),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 9),
        ("ALIGN",       (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [_rl_color(COLORS["white"]), _rl_color(COLORS["light"])]),
        ("GRID",        (0, 0), (-1, -1), 0.5, _rl_color(COLORS["mid"])),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",  (0, 0), (-1, -1), 6),
    ])
    # Colour the score cell
    risk_overall = p["narrative"].get("overall_risk_rating", "")
    rc = RISK_COLOR.get(risk_overall, COLORS["mid"])
    ts_style.add("BACKGROUND", (1, 3), (1, 3), _rl_color(rc))
    ts_style.add("TEXTCOLOR",  (1, 3), (1, 3), colors.white)
    ts_style.add("FONTNAME",   (1, 3), (1, 3), "Helvetica-Bold")
    t.setStyle(ts_style)
    return t


def _findings_table(findings: list, styles: dict, status_filter=None) -> "Table":
    """Findings table for a given status filter."""
    if status_filter:
        rows = [f for f in findings if f["status"] in status_filter]
    else:
        rows = findings

    col_widths = [1.5*cm, 5.5*cm, 1.0*cm, 1.5*cm, 1.8*cm, 6.5*cm]
    header = [
        Paragraph("ID", styles["small_white"]),
        Paragraph("Description", styles["small_white"]),
        Paragraph("Lvl", styles["small_white"]),
        Paragraph("Status", styles["small_white"]),
        Paragraph("Risk", styles["small_white"]),
        Paragraph("Remediation", styles["small_white"]),
    ]
    data = [header]
    ts_cmds = [
        ("BACKGROUND",  (0, 0), (-1, 0), _rl_color(COLORS["header"])),
        ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
        ("FONTSIZE",    (0, 0), (-1, -1), 8),
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("GRID",        (0, 0), (-1, -1), 0.3, _rl_color(COLORS["mid"])),
        ("TOPPADDING",  (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
    ]

    for i, f in enumerate(rows, start=1):
        ai     = f.get("ai_analysis", {})
        status = f.get("status", "")
        risk   = ai.get("risk_level", "")
        s_c    = STATUS_COLOR.get(status, COLORS["mid"])
        r_c    = RISK_COLOR.get(risk, COLORS["mid"])
        rem    = ai.get("remediation_steps", f.get("remediation", ""))
        if isinstance(rem, list):
            rem = "; ".join(rem[:3])
        bg = COLORS["light"] if i % 2 == 0 else COLORS["white"]

        if not risk:
            risk = {"PASS": "Passed", "SKIPPED": "N/A", "ERROR": "Error"}.get(status, "Unknown")
            r_c  = COLORS["mid"]
        row = [
            Paragraph(_esc(f.get("control_id", "")), styles["small"]),
            Paragraph(_esc(f.get("description", "")[:90]), styles["small"]),
            Paragraph(_esc(f.get("level", "")), styles["small"]),
            Paragraph(f"<b>{_esc(status)}</b>", styles["small"]),
            Paragraph(f"<b>{_esc(risk)}</b>", styles["small"]),
            Paragraph(_esc(str(rem)[:200]), styles["small"]),
        ]
        data.append(row)
        ts_cmds.append(("BACKGROUND", (0, i), (-1, i), _rl_color(bg)))
        ts_cmds.append(("BACKGROUND", (3, i), (3, i), _rl_color(s_c)))
        ts_cmds.append(("TEXTCOLOR",  (3, i), (3, i), colors.white))
        ts_cmds.append(("BACKGROUND", (4, i), (4, i), _rl_color(r_c)))
        ts_cmds.append(("TEXTCOLOR",  (4, i), (4, i), colors.white))

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle(ts_cmds))
    return t


def _chain_table(chains: list, styles: dict) -> "Table":
    col_widths = [1.5*cm, 3.5*cm, 2*cm, 3.5*cm, 7.5*cm]
    header = [
        Paragraph("ID", styles["small_white"]),
        Paragraph("Chain Name", styles["small_white"]),
        Paragraph("Risk", styles["small_white"]),
        Paragraph("Controls", styles["small_white"]),
        Paragraph("Attack Narrative", styles["small_white"]),
    ]
    data   = [header]
    ts_cmds = [
        ("BACKGROUND",  (0, 0), (-1, 0), _rl_color(COLORS["header"])),
        ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
        ("FONTSIZE",    (0, 0), (-1, -1), 8),
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
        ("GRID",        (0, 0), (-1, -1), 0.3, _rl_color(COLORS["mid"])),
        ("TOPPADDING",  (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
    ]
    for i, chain in enumerate(chains, start=1):
        risk     = chain.get("risk_level", "")
        r_c      = RISK_COLOR.get(risk, COLORS["mid"])
        controls = chain.get("controls_involved", [])
        if isinstance(controls, list):
            controls = ", ".join(controls)
        row = [
            Paragraph(chain.get("chain_id", ""), styles["small"]),
            Paragraph(chain.get("chain_name", ""), styles["small"]),
            Paragraph(f"<b>{risk}</b>", styles["small"]),
            Paragraph(controls, styles["small"]),
            Paragraph(chain.get("attack_narrative", "")[:200], styles["small"]),
        ]
        data.append(row)
        bg = COLORS["light"] if i % 2 == 0 else COLORS["white"]
        ts_cmds.append(("BACKGROUND", (0, i), (-1, i), _rl_color(bg)))
        ts_cmds.append(("BACKGROUND", (2, i), (2, i), _rl_color(r_c)))
        ts_cmds.append(("TEXTCOLOR",  (2, i), (2, i), colors.white))

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle(ts_cmds))
    return t


def _compliance_table(p: dict, styles: dict) -> "Table":
    """Build a ReportLab Table showing CIS → framework control mappings."""
    from tools.compliance_mappings import FRAMEWORKS

    frameworks = p.get("frameworks", [])
    fw_names   = [FRAMEWORKS[fw]["name"] for fw in frameworks if fw in FRAMEWORKS]

    headers = ["ID", "Description"] + fw_names
    n_fw    = len(fw_names)
    # Widths: ID=1.2cm, Desc shares remaining space with fw cols
    id_w   = 1.2 * cm
    desc_w = 5.5 * cm
    fw_w   = (17 * cm - id_w - desc_w) / max(n_fw, 1)
    col_widths = [id_w, desc_w] + [fw_w] * n_fw

    hdr_style = ParagraphStyle("ch", fontSize=8, fontName="Helvetica-Bold",
                                textColor=colors.white, alignment=1)
    id_style  = ParagraphStyle("ci", fontSize=7, fontName="Helvetica-Bold", alignment=1)
    txt_style = ParagraphStyle("ct", fontSize=7, fontName="Helvetica", leading=9)
    fw_colors_hex = {
        "nist_csf":     "#1A5276",
        "iso_27001":    "#1E8449",
        "pci_dss":      "#7D3C98",
        "mitre_attack": "#922B21",
    }

    rows = [[Paragraph(h, hdr_style) for h in headers]]
    ts_cmds = [
        ("BACKGROUND",  (0, 0), (-1, 0), _rl_color(COLORS["header"])),
        ("GRID",        (0, 0), (-1, -1), 0.4, colors.Color(0.75, 0.75, 0.75)),
        ("TOPPADDING",  (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN",      (0, 0), (-1, -1), "TOP"),
    ]

    for i, finding in enumerate(p["findings"], start=1):
        compliance = finding.get("compliance", {})
        status     = finding.get("status", "")
        bg         = _rl_color(COLORS["light"]) if i % 2 == 0 else colors.white
        s_color    = _rl_color(STATUS_COLOR.get(status, COLORS["mid"]))

        id_para   = Paragraph(f"<b>{_esc(finding.get('control_id',''))}</b>", id_style)
        desc_para = Paragraph(_esc(finding.get("description", "")), txt_style)

        fw_cells = []
        for fw in frameworks:
            controls = compliance.get(fw, [])
            text     = "<br/>".join(_esc(c) for c in controls) if controls else "<font color='#AAAAAA'>—</font>"
            fw_hex   = fw_colors_hex.get(fw, "#333333")
            cell_style = ParagraphStyle("fw", fontSize=7, fontName="Helvetica",
                                        textColor=_rl_color(fw_hex) if controls else _rl_color("#AAAAAA"),
                                        leading=9)
            fw_cells.append(Paragraph(text, cell_style))

        rows.append([id_para, desc_para] + fw_cells)
        ts_cmds.append(("BACKGROUND", (0, i), (-1, i), bg))
        ts_cmds.append(("BACKGROUND", (0, i), (0, i), s_color))
        ts_cmds.append(("TEXTCOLOR",  (0, i), (0, i), colors.white))

    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle(ts_cmds))
    return t


def _page_header_footer(canvas, doc):
    canvas.saveState()
    w, h = A4
    # Header bar
    canvas.setFillColor(_rl_color(COLORS["header"]))
    canvas.rect(0, h - 1.4*cm, w, 1.4*cm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(1*cm, h - 0.9*cm, "FW AI Audit — Security Assessment Report")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(w - 1*cm, h - 0.9*cm, f"CONFIDENTIAL")
    # Footer bar
    canvas.setFillColor(_rl_color(COLORS["light"]))
    canvas.rect(0, 0, w, 0.9*cm, fill=1, stroke=0)
    canvas.setFillColor(_rl_color(COLORS["header"]))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(1*cm, 0.3*cm, f"Page {doc.page}")
    canvas.drawCentredString(w/2, 0.3*cm, "Firewall Security Audit")
    canvas.restoreState()


def generate_pdf(data: dict, output_path: str) -> str:
    if not HAS_REPORTLAB:
        raise RuntimeError("reportlab required: pip install reportlab --break-system-packages")

    p  = _prepare(data)
    st = _build_styles()

    doc = _CISDocTemplate(
        output_path,
        pagesize=A4,
        topMargin=1.8*cm, bottomMargin=1.4*cm,
        leftMargin=1.6*cm, rightMargin=1.6*cm,
        title="FW AI Audit Security Assessment",
        author="AI Audit Engine",
    )
    doc.set_header_footer(_page_header_footer)

    story = []
    meta  = p["meta"]
    narr  = p["narrative"]

    # ---- Cover Page --------------------------------------------------------
    story.append(Spacer(1, 1.5*cm))

    # Cover banner
    benchmark = meta.get("benchmark", "FW AI Audit Security Assessment")
    cover_data = [[_para(benchmark.upper(), st["title"])],
                  [_para("Security Assessment Report", st["subtitle"])]]
    cover_table = Table(cover_data, colWidths=[17*cm])
    cover_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _rl_color(COLORS["header"])),
        ("TOPPADDING",    (0, 0), (-1, -1), 14),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
    ]))
    story.append(cover_table)
    story.append(Spacer(1, 0.8*cm))

    # Target info table
    target  = meta.get("target", "Unknown")
    ts      = meta.get("generated", "")[:19].replace("T", " ") + " UTC"
    ctx     = meta.get("device_context", {})
    org     = ctx.get("organization", "")
    role    = ctx.get("device_role", "")
    industry = ctx.get("industry", "")

    info_data = [
        ["Target Device", target],
        ["Assessment Date", ts],
        ["Organization", org or "—"],
        ["Device Role", role or "—"],
        ["Industry", industry or "—"],
        ["Framework", meta.get("benchmark", "FW AI Audit Security Assessment")],
        ["Classification", "CONFIDENTIAL"],
    ]
    info_table = Table(info_data, colWidths=[5*cm, 12*cm])
    info_table.setStyle(TableStyle([
        ("FONTNAME",  (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",  (0, 0), (-1, -1), 9.5),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1),
         [_rl_color(COLORS["light"]), _rl_color(COLORS["white"])]),
        ("GRID",      (0, 0), (-1, -1), 0.5, _rl_color(COLORS["mid"])),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.6*cm))

    # Compliance score + stats
    overall = narr.get("overall_risk_rating", "Unknown")
    score_data = [[
        Paragraph(f"<b>{p['score']}%</b>", ParagraphStyle("sc",
            fontSize=32, fontName="Helvetica-Bold",
            textColor=_rl_color(COLORS["header"]), alignment=TA_CENTER)),
        Paragraph(f"<b>Overall Risk: {_esc(overall)}</b>", ParagraphStyle("or",
            fontSize=13, fontName="Helvetica-Bold",
            textColor=_rl_color(COLORS["white"]), alignment=TA_CENTER,
            backColor=_rl_color(RISK_COLOR.get(overall, COLORS["mid"])))),
        _stat_table(p, st),
    ]]
    score_table = Table(score_data, colWidths=[4*cm, 5*cm, 8*cm])
    score_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(score_table)
    story.append(PageBreak())

    # ---- Table of Contents ------------------------------------------------
    toc = TableOfContents()
    toc.levelStyles = [st["toc_h1"], st["toc_h2"]]
    toc.dotsMinLevel = 0
    story.append(_para("Table of Contents", st["section_header"]))
    story.append(Spacer(1, 0.3*cm))
    story.append(toc)
    story.append(PageBreak())

    # ---- Executive Summary -------------------------------------------------
    story.append(_para("Executive Summary", st["section_header"]))
    story.append(Spacer(1, 0.3*cm))
    exec_s = narr.get("executive_summary", {})
    if isinstance(exec_s, dict):
        headline = exec_s.get("headline", "")
        if headline:
            story.append(Paragraph(f"<i>{_esc(headline)}</i>", st["body_bold"]))
            story.append(Spacer(1, 0.2*cm))
        for key in ["paragraph_1", "paragraph_2", "paragraph_3"]:
            txt = exec_s.get(key, "")
            if txt:
                story.append(_para(txt, st["body"]))
    elif isinstance(exec_s, str):
        story.append(_para(exec_s, st["body"]))

    story.append(Spacer(1, 0.4*cm))

    # Compliance score interpretation
    interp = narr.get("compliance_score_interpretation", "")
    if interp:
        story.append(_para("Score Interpretation", st["heading2"]))
        story.append(_para(interp, st["body"]))
        story.append(Spacer(1, 0.3*cm))

    # Positive findings
    positive = narr.get("positive_findings", "")
    if positive:
        story.append(_para("Positive Security Findings", st["heading2"]))
        story.extend(_multi_para(positive, st["body"]))
        story.append(Spacer(1, 0.3*cm))

    # Limitations
    limitations = narr.get("assessment_limitations", "")
    if limitations:
        story.append(_para("Assessment Limitations", st["heading2"]))
        story.extend(_multi_para(limitations, st["body"]))

    story.append(Spacer(1, 0.4*cm))

    # ---- Technical Summary -------------------------------------------------
    story.append(_para("Technical Summary", st["section_header"]))
    tech_s = narr.get("technical_summary", {})
    if isinstance(tech_s, dict):
        for key in ["paragraph_1", "paragraph_2"]:
            txt = tech_s.get(key, "")
            if txt:
                story.extend(_multi_para(txt, st["body"]))
                story.append(Spacer(1, 0.2*cm))
    elif isinstance(tech_s, str):
        story.extend(_multi_para(tech_s, st["body"]))

    story.append(Spacer(1, 0.4*cm))

    # Top 5 actions
    story.append(_para("Priority Remediation Actions", st["heading2"]))
    actions = narr.get("top_5_priority_actions", [])
    if actions:
        act_data = [[
            Paragraph("<b>#</b>", st["small"]),
            Paragraph("<b>Action</b>", st["small"]),
            Paragraph("<b>Justification</b>", st["small"]),
            Paragraph("<b>Effort</b>", st["small"]),
            Paragraph("<b>Impact</b>", st["small"]),
        ]]
        for action in actions:
            effort  = action.get("effort", "")
            e_color = {"Low": COLORS["pass"], "Medium": COLORS["manual"],
                       "High": COLORS["fail"]}.get(effort, COLORS["mid"])
            act_data.append([
                Paragraph(str(action.get("rank", "")), st["small"]),
                Paragraph(str(action.get("action", "")), st["small"]),
                Paragraph(str(action.get("justification", "")), st["small"]),
                Paragraph(f"<b>{effort}</b>", st["small"]),
                Paragraph(str(action.get("impact", "")), st["small"]),
            ])
        act_table = Table(act_data, colWidths=[0.8*cm, 4*cm, 5*cm, 1.8*cm, 5.4*cm])
        ts_act = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), _rl_color(COLORS["accent"])),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
            ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ("VALIGN",     (0, 0), (-1, -1), "TOP"),
            ("GRID",       (0, 0), (-1, -1), 0.3, _rl_color(COLORS["mid"])),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [_rl_color(COLORS["white"]), _rl_color(COLORS["light"])]),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ])
        for i, action in enumerate(actions, start=1):
            effort  = action.get("effort", "")
            e_color = {"Low": COLORS["pass"], "Medium": COLORS["manual"],
                       "High": COLORS["fail"]}.get(effort, COLORS["mid"])
            ts_act.add("BACKGROUND", (3, i), (3, i), _rl_color(e_color))
            ts_act.add("TEXTCOLOR",  (3, i), (3, i), colors.white)
        act_table.setStyle(ts_act)
        story.append(act_table)

    story.append(PageBreak())

    # ---- Attack Chains -----------------------------------------------------
    if p["chains"]:
        story.append(_para("Attack Chain Analysis", st["section_header"]))
        story.append(Spacer(1, 0.2*cm))
        story.append(_para(
            "The following compound risk scenarios were identified where multiple "
            "control failures combine to create threats significantly worse than "
            "any single finding in isolation.", st["body"]))
        story.append(Spacer(1, 0.3*cm))

        for chain in p["chains"]:
            risk       = chain.get("risk_level", "")
            r_c        = RISK_COLOR.get(risk, COLORS["mid"])
            chain_id   = _esc(chain.get("chain_id", ""))
            chain_name = _esc(chain.get("chain_name", ""))
            controls   = ", ".join(chain.get("controls_involved", []))
            narrative  = chain.get("attack_narrative", "")
            blast      = chain.get("blast_radius", "")
            fix_order  = _esc(" → ".join(chain.get("priority_fix_order", [])))
            chain_block = [
                # Chain title bar
                Table([[
                    Paragraph(f"<b>{chain_id} — {chain_name}</b>", st["body_bold"]),
                    Paragraph(f"<b>{_esc(risk)}</b>", ParagraphStyle("rl",
                        fontSize=9, fontName="Helvetica-Bold",
                        textColor=_rl_color(COLORS["white"]),
                        backColor=_rl_color(r_c), alignment=TA_CENTER,
                        leading=14)),
                ]], colWidths=[13*cm, 4*cm]),
                Spacer(1, 0.25*cm),
                # Controls involved
                Paragraph("<b>Controls Involved</b>", st["field_label"]),
                Paragraph(_esc(controls), st["field_value"]),
                Spacer(1, 0.2*cm),
                # Attack narrative — each numbered step as its own paragraph
                Paragraph("<b>Attack Narrative</b>", st["field_label"]),
            ] + _multi_para(narrative, st["field_value"], gap=0.1) + [
                Spacer(1, 0.2*cm),
                # Blast radius
                Paragraph("<b>Blast Radius</b>", st["field_label"]),
            ] + _multi_para(blast, st["field_value"], gap=0.1) + [
                Spacer(1, 0.2*cm),
                # Fix order
                Paragraph("<b>Recommended Fix Order</b>", st["field_label"]),
                Paragraph(fix_order, st["field_value"]),
                Spacer(1, 0.35*cm),
                HRFlowable(width="100%", thickness=0.5, color=_rl_color(COLORS["mid"])),
                Spacer(1, 0.3*cm),
            ]
            story.extend(chain_block)

        story.append(PageBreak())

    # ---- FAIL Findings Details ---------------------------------------------
    fails = [f for f in p["findings"] if f["status"] == "FAIL"]
    if fails:
        story.append(_para(f"Automated Failures ({len(fails)} findings)", st["section_header"]))
        story.append(Spacer(1, 0.2*cm))
        story.append(_findings_table(p["findings"], st, status_filter=["FAIL"]))
        story.append(PageBreak())

    # ---- MANUAL Findings ---------------------------------------------------
    manuals = [f for f in p["findings"] if f["status"] == "MANUAL"]
    if manuals:
        story.append(_para(f"Manual Review Required ({len(manuals)} findings)", st["section_header"]))
        story.append(Spacer(1, 0.2*cm))
        story.append(_para(
            "The following controls cannot be verified through automated means. "
            "Each requires manual verification in SmartConsole or via expert review.",
            st["body"]))
        story.append(Spacer(1, 0.2*cm))
        story.append(_findings_table(p["findings"], st, status_filter=["MANUAL"]))
        story.append(Spacer(1, 0.4*cm))

    # ---- All Findings (Appendix) -------------------------------------------
    story.append(_para("Appendix: Complete Findings", st["section_header"]))
    story.append(Spacer(1, 0.2*cm))
    story.append(_findings_table(p["findings"], st))

    # ---- Compliance Mapping (Appendix) — only when frameworks selected -----
    if p.get("frameworks"):
        story.append(PageBreak())
        story.append(_para("Appendix: Security Compliance Mapping", st["section_header"]))
        story.append(Spacer(1, 0.3*cm))
        story.append(_compliance_table(p, st))

    doc.multiBuild(story)
    print(f"[REPORT] ✓ PDF  → {os.path.abspath(output_path)}")
    return output_path


# ============================================================================
# 4. WORD REPORT (.docx)
# ============================================================================
try:
    from docx import Document as _DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _docx_set_cell_bg(cell, hex_color: str):
    hex_color = hex_color.lstrip("#")
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = _OxmlElement("w:shd")
    shd.set(_qn("w:val"),   "clear")
    shd.set(_qn("w:color"), "auto")
    shd.set(_qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _docx_cell_text(cell, text: str, bold=False, italic=False,
                    color_hex: str = None, size_pt: int = 10,
                    align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(_strip_html(str(text)))
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)


def _docx_add_toc(doc):
    """Insert a Word TOC field (auto-populated when opened in Word)."""
    para = doc.add_paragraph()
    run  = para.add_run()
    fld  = _OxmlElement("w:fldChar")
    fld.set(_qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run2   = para.add_run()
    instr  = _OxmlElement("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instr)

    run3 = para.add_run()
    fld2 = _OxmlElement("w:fldChar")
    fld2.set(_qn("w:fldCharType"), "separate")
    run3._r.append(fld2)

    run4 = para.add_run()
    t    = _OxmlElement("w:t")
    t.text = "[Right-click and select Update Field to refresh TOC]"
    run4._r.append(t)

    run5 = para.add_run()
    fld3 = _OxmlElement("w:fldChar")
    fld3.set(_qn("w:fldCharType"), "end")
    run5._r.append(fld3)


def _docx_findings_table(doc, findings: list, status_filter=None):
    rows = [f for f in findings if f["status"] in status_filter] \
           if status_filter else findings
    if not rows:
        return

    headers = ["ID", "Description", "Level", "Status", "Risk", "Remediation"]
    widths  = [Cm(1.8), Cm(5.5), Cm(1.4), Cm(1.8), Cm(2.2), Cm(6)]
    table   = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        hdr.cells[i].width = w
        _docx_set_cell_bg(hdr.cells[i], COLORS["header"])
        _docx_cell_text(hdr.cells[i], h, bold=True, color_hex=COLORS["white"],
                        size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, f in enumerate(rows, start=1):
        ai     = f.get("ai_analysis", {})
        status = f.get("status", "")
        risk   = ai.get("risk_level", "")
        rem    = ai.get("remediation_steps", f.get("remediation", ""))
        if isinstance(rem, list):
            rem = "; ".join(rem[:3])
        s_hex  = STATUS_COLOR.get(status, COLORS["mid"])
        r_hex  = RISK_COLOR.get(risk,   COLORS["mid"])
        bg_hex = COLORS["light"] if ri % 2 == 0 else COLORS["white"]

        cells = table.rows[ri].cells
        for i, w in enumerate(widths):
            cells[i].width = w

        _docx_cell_text(cells[0], f.get("control_id", ""), size_pt=8)
        _docx_cell_text(cells[1], f.get("description", ""), size_pt=8)
        _docx_cell_text(cells[2], f.get("level", ""),       size_pt=8,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        _docx_set_cell_bg(cells[3], s_hex)
        _docx_cell_text(cells[3], status, bold=True, color_hex=COLORS["white"],
                        size_pt=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        if risk:
            _docx_set_cell_bg(cells[4], r_hex)
            _docx_cell_text(cells[4], risk, bold=True, color_hex=COLORS["white"],
                            size_pt=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            _docx_cell_text(cells[4], "—", size_pt=8,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        _docx_cell_text(cells[5], str(rem)[:220], size_pt=8)

        for i in range(len(headers)):
            if i not in (3, 4):
                _docx_set_cell_bg(cells[i], bg_hex)


def _docx_compliance_table(doc, p: dict):
    """Add a compliance mapping table to a docx document."""
    from tools.compliance_mappings import FRAMEWORKS

    frameworks = p.get("frameworks", [])
    if not frameworks:
        return

    fw_names = [FRAMEWORKS[fw]["name"] for fw in frameworks if fw in FRAMEWORKS]
    headers  = ["ID", "Description"] + fw_names
    n_cols   = len(headers)

    table = doc.add_table(rows=1 + len(p["findings"]), cols=n_cols)
    table.style = "Table Grid"

    col_widths = [Cm(1.6), Cm(5.5)] + [Cm(18.0 / max(len(frameworks), 1) - 0.4)] * len(frameworks)

    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr.cells[i].width = w
        _docx_set_cell_bg(hdr.cells[i], COLORS["header"])
        _docx_cell_text(hdr.cells[i], h, bold=True, color_hex=COLORS["white"],
                        size_pt=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    fw_colors = {
        "nist_csf":     "1A5276",
        "iso_27001":    "1E8449",
        "pci_dss":      "7D3C98",
        "mitre_attack": "922B21",
    }

    for ri, finding in enumerate(p["findings"], start=1):
        compliance = finding.get("compliance", {})
        status     = finding.get("status", "")
        bg_hex     = COLORS["light"] if ri % 2 == 0 else COLORS["white"]
        s_hex      = STATUS_COLOR.get(status, COLORS["mid"])
        cells      = table.rows[ri].cells

        for i, w in enumerate(col_widths):
            cells[i].width = w

        _docx_set_cell_bg(cells[0], s_hex)
        _docx_cell_text(cells[0], finding.get("control_id", ""), bold=True,
                        color_hex=COLORS["white"], size_pt=8,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

        _docx_set_cell_bg(cells[1], bg_hex)
        _docx_cell_text(cells[1], finding.get("description", ""), size_pt=8)

        for fi, fw in enumerate(frameworks, start=2):
            controls = compliance.get(fw, [])
            text     = "  ".join(controls) if controls else "—"
            color    = fw_colors.get(fw, "333333") if controls else "AAAAAA"
            _docx_set_cell_bg(cells[fi], bg_hex)
            _docx_cell_text(cells[fi], text, size_pt=7, color_hex=color)


def generate_docx(data: dict, output_path: str) -> str:
    if not HAS_DOCX:
        raise RuntimeError("python-docx required: uv add python-docx")

    p    = _prepare(data)
    meta = p["meta"]
    narr = p["narrative"]
    doc  = _DocxDocument()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def h1(text):
        doc.add_heading(text, level=1)

    def h2(text):
        doc.add_heading(text, level=2)

    def body(text):
        doc.add_paragraph(_strip_html(str(text)))

    def multi_body(text):
        for chunk in _split_numbered(str(text)):
            doc.add_paragraph(_strip_html(chunk))

    # ── Cover ──────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("FW AI Audit — Firewall Security Assessment")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x25, 0x2F)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run("Security Assessment Report")
    run2.font.size = Pt(14)
    run2.italic = True

    doc.add_paragraph()

    ctx    = meta.get("device_context", {})
    target = meta.get("target", "Unknown")
    ts     = meta.get("generated", "")[:19].replace("T", " ") + " UTC"
    overall = narr.get("overall_risk_rating", "Unknown")
    score   = p["score"]

    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = "Table Grid"
    info_rows = [
        ("Target Device",   target),
        ("Assessment Date", ts),
        ("Organization",    ctx.get("organization", "—")),
        ("Device Role",     ctx.get("device_role", "—")),
        ("Industry",        ctx.get("industry", "—")),
        ("Framework",       meta.get("benchmark", "FW AI Audit Security Assessment")),
        ("Classification",  "CONFIDENTIAL"),
        ("Overall Risk",    f"{overall}  |  Score: {score}%"),
    ]
    for ri, (label, value) in enumerate(info_rows):
        cells = info_table.rows[ri].cells
        _docx_set_cell_bg(cells[0], COLORS["light"])
        _docx_cell_text(cells[0], label, bold=True, size_pt=10)
        _docx_cell_text(cells[1], value, size_pt=10)

    # Colour the Overall Risk row
    risk_hex = RISK_COLOR.get(overall, COLORS["mid"])
    _docx_set_cell_bg(info_table.rows[7].cells[1], risk_hex)
    _docx_cell_text(info_table.rows[7].cells[1], f"{overall}  |  Score: {score}%",
                    bold=True, color_hex=COLORS["white"], size_pt=10)

    doc.add_paragraph()

    # Score summary mini-table
    stat_table = doc.add_table(rows=2, cols=4)
    stat_table.style = "Table Grid"
    stat_headers = ["Pass", "Fail", "Manual", "Skipped"]
    stat_values  = [p["passes"], p["fails"], p["manual"], p["skipped"]]
    stat_colors  = [COLORS["pass"], COLORS["fail"], COLORS["manual"], COLORS["skipped"]]
    for ci, (h, v, c) in enumerate(zip(stat_headers, stat_values, stat_colors)):
        _docx_set_cell_bg(stat_table.rows[0].cells[ci], c)
        _docx_cell_text(stat_table.rows[0].cells[ci], h, bold=True,
                        color_hex=COLORS["white"], size_pt=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        _docx_set_cell_bg(stat_table.rows[1].cells[ci], c)
        _docx_cell_text(stat_table.rows[1].cells[ci], str(v), bold=True,
                        color_hex=COLORS["white"], size_pt=14,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────────────
    h1("Table of Contents")
    _docx_add_toc(doc)
    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────
    h1("Executive Summary")
    exec_s = narr.get("executive_summary", {})
    if isinstance(exec_s, dict):
        headline = exec_s.get("headline", "")
        if headline:
            p_hl = doc.add_paragraph()
            run  = p_hl.add_run(_strip_html(headline))
            run.bold   = True
            run.italic = True
        for key in ["paragraph_1", "paragraph_2", "paragraph_3"]:
            txt = exec_s.get(key, "")
            if txt:
                multi_body(txt)
    elif isinstance(exec_s, str):
        multi_body(exec_s)

    h2("Score Interpretation")
    interp = narr.get("compliance_score_interpretation", "")
    if interp:
        multi_body(interp)

    h2("Positive Security Findings")
    positive = narr.get("positive_findings", "")
    if positive:
        multi_body(positive)

    h2("Assessment Limitations")
    limitations = narr.get("assessment_limitations", "")
    if limitations:
        multi_body(limitations)

    # ── Technical Summary ─────────────────────────────────────────────────
    h1("Technical Summary")
    tech_s = narr.get("technical_summary", {})
    if isinstance(tech_s, dict):
        for key in ["paragraph_1", "paragraph_2"]:
            txt = tech_s.get(key, "")
            if txt:
                multi_body(txt)
    elif isinstance(tech_s, str):
        multi_body(tech_s)

    h2("Priority Remediation Actions")
    actions = narr.get("top_5_priority_actions", [])
    if actions:
        act_table = doc.add_table(rows=1 + len(actions), cols=5)
        act_table.style = "Table Grid"
        act_headers = ["#", "Action", "Justification", "Effort", "Impact"]
        hrow = act_table.rows[0]
        for ci, h_txt in enumerate(act_headers):
            _docx_set_cell_bg(hrow.cells[ci], COLORS["accent"])
            _docx_cell_text(hrow.cells[ci], h_txt, bold=True,
                            color_hex=COLORS["white"], size_pt=9,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        for ri, action in enumerate(actions, start=1):
            effort   = action.get("effort", "")
            e_hex    = {"Low": COLORS["pass"], "Medium": COLORS["manual"],
                        "High": COLORS["fail"]}.get(effort, COLORS["mid"])
            bg_hex   = COLORS["light"] if ri % 2 == 0 else COLORS["white"]
            cells    = act_table.rows[ri].cells
            vals     = [str(action.get("rank", "")), action.get("action", ""),
                        action.get("justification", ""), effort,
                        action.get("impact", "")]
            for ci, val in enumerate(vals):
                _docx_set_cell_bg(cells[ci], e_hex if ci == 3 else bg_hex)
                _docx_cell_text(cells[ci], val, size_pt=9,
                                bold=(ci == 3),
                                color_hex=COLORS["white"] if ci == 3 else None)

    # ── Attack Chain Analysis ─────────────────────────────────────────────
    if p["chains"]:
        doc.add_page_break()
        h1("Attack Chain Analysis")
        body("The following compound risk scenarios were identified where multiple "
             "control failures combine to create threats significantly worse than "
             "any single finding in isolation.")
        doc.add_paragraph()

        for chain in p["chains"]:
            risk     = chain.get("risk_level", "")
            r_hex    = RISK_COLOR.get(risk, COLORS["mid"])
            controls = ", ".join(chain.get("controls_involved", []))
            narrative_txt = chain.get("attack_narrative", "")
            blast_txt     = chain.get("blast_radius", "")
            fix_order     = " → ".join(chain.get("priority_fix_order", []))
            chain_id   = chain.get("chain_id", "")
            chain_name = chain.get("chain_name", "")

            # Chain title
            ct = doc.add_paragraph()
            r1 = ct.add_run(f"{chain_id} — {chain_name}")
            r1.bold = True
            r1.font.size = Pt(11)
            r1.font.color.rgb = RGBColor(0x1A, 0x25, 0x2F)
            r2 = ct.add_run(f"  [{risk}]")
            r2.bold = True
            r2.font.size = Pt(10)
            risk_r, risk_g, risk_b = hex_to_rgb(r_hex)
            r2.font.color.rgb = RGBColor(risk_r, risk_g, risk_b)

            # Fields
            for label, text in [
                ("Controls Involved", controls),
                ("Attack Narrative",  narrative_txt),
                ("Blast Radius",      blast_txt),
                ("Recommended Fix Order", fix_order),
            ]:
                lp = doc.add_paragraph()
                lr = lp.add_run(label)
                lr.bold = True
                lr.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)  # accent blue
                if label == "Attack Narrative":
                    for chunk in _split_numbered(text):
                        doc.add_paragraph(_strip_html(chunk), style="List Number")
                else:
                    doc.add_paragraph(_strip_html(text))

            doc.add_paragraph("─" * 80)

    # ── Automated Failures ────────────────────────────────────────────────
    fails = [f for f in p["findings"] if f["status"] == "FAIL"]
    if fails:
        doc.add_page_break()
        h1(f"Automated Failures ({len(fails)} findings)")
        _docx_findings_table(doc, p["findings"], status_filter=["FAIL"])

    # ── Appendix: Complete Findings ───────────────────────────────────────
    doc.add_page_break()
    h1("Appendix: Complete Findings")
    _docx_findings_table(doc, p["findings"])

    # ── Appendix: Security Compliance Mapping ─────────────────────────────
    if p.get("frameworks"):
        doc.add_page_break()
        h1("Appendix: Security Compliance Mapping")
        _docx_compliance_table(doc, p)

    doc.save(output_path)
    print(f"[REPORT] ✓ Word → {os.path.abspath(output_path)}")
    return output_path


# ============================================================================
# Convenience: generate all formats at once
# ============================================================================
# ============================================================================
# 5. CONSULTING REPORT (.docx) — follows data-source/template_1.txt structure
# ============================================================================

# Section-to-check-ID mapping.  Each tuple: (section_num, title, [(sub_num, sub_title, [ids])])
_CONSULTING_MAP = [
    ("2", "Security Policy Review", [
        ("2.1", "Rules with Zero Hit Count",       ["RQ-2"]),
        ("2.2", "Overly Permissive Rules",          ["3.5", "3.6", "3.7", "RQ-3"]),
        ("2.3", "Shadow Rules",                     ["RQ-1"]),
    ]),
    ("3", "NAT Policy Review", [
        ("3.1", "NAT Rules with Zero Hit Count",    ["NAT-1"]),
        ("3.2", "Any-Any NAT Rules",                ["NAT-2"]),
    ]),
    ("4", "Firewall Posture Assessment", [
        ("4.1", "Stealth and Cleanup Rules",        ["3.1", "3.2", "3.3"]),
        ("4.2", "Logging and Visibility",           ["3.8", "2.6.1", "2.6.2"]),
        ("4.3", "Anti-Spoofing and Stateful Inspection", ["3.10", "3.11", "3.12"]),
    ]),
    ("5", "Software Version and Lifecycle Risk", []),   # special: sourced from meta
    ("6", "Additional Security, Architecture, and Governance Findings", [
        ("6.1",  "Password Policy",                 ["1.1", "1.2", "1.3",
                                                     "1.4a", "1.4b",
                                                     "1.5", "1.6", "1.7"]),
        ("6.2",  "Account Lockout Controls",        ["1.11", "1.12", "1.13"]),
        ("6.3",  "Inactive and Non-Expiring Accounts", ["1.8", "1.9", "1.10"]),
        ("6.4",  "Admin Privilege and Identity Management", ["IAM-1", "IAM-2", "IAM-3"]),
        ("6.5",  "Break-Glass Administrator Model", ["GOV-3"]),
        ("6.6",  "Session Timeout and Access Control", ["2.5.1", "2.5.2",
                                                        "2.5.4", "2.5.5"]),
        ("6.7",  "Login Banner and MOTD",           ["2.1.1", "2.1.2"]),
        ("6.8",  "SNMP Security",                   ["2.2.1", "2.2.2", "2.2.3", "2.2.4"]),
        ("6.9",  "Backup and Recovery",             ["2.4.1", "2.4.2", "2.4.3"]),
        ("6.10", "Time Synchronisation",            ["2.3.1", "2.3.2"]),
        ("6.11", "General Hardening",               ["2.1.3", "2.1.4", "2.1.5", "2.1.6",
                                                     "2.1.7", "2.1.8", "2.1.9", "2.1.10",
                                                     "3.14", "3.15", "3.16",
                                                     "3.17", "3.18", "3.19"]),
        ("6.12", "Firewall Deployment Architecture", ["GOV-1"]),
        ("6.13", "Failover Testing Plan",           ["GOV-2"]),
        ("6.14", "Change Management and Governance", ["GOV-4"]),
        ("6.15", "Hit Count Tracking and Audit",    ["3.4"]),
    ]),
]

_PANOS_CONSULTING_MAP = [
    ("2", "Security Policy Review", [
        ("2.1", "Application and Service Policy Controls",      ["7.1", "7.2"]),
        ("2.2", "Threat Intelligence-Based Blocking",           ["7.3"]),
        ("2.3", "Default Security Policy Logging",              ["7.4"]),
    ]),
    ("3", "Threat Prevention and Security Profiles", [
        ("3.1", "Security Profile Coverage",                    ["6.1", "6.2", "6.3", "6.4",
                                                                  "6.5", "6.6", "6.7"]),
        ("3.2", "URL Filtering and Content Inspection",         ["6.8", "6.9", "6.10", "6.11",
                                                                  "6.12", "6.13", "6.14",
                                                                  "6.19", "6.20", "6.21",
                                                                  "6.22", "6.23", "6.24", "6.25"]),
        ("3.3", "WildFire Malware Analysis",                    ["5.1", "5.2", "5.3", "5.4",
                                                                  "5.5", "5.6", "5.7"]),
    ]),
    ("4", "Firewall Posture Assessment", [
        ("4.1", "High Availability and Resilience",             ["3.1", "3.2", "3.3"]),
        ("4.2", "Zone Protection Controls",                     ["6.15", "6.16", "6.17", "6.18"]),
        ("4.3", "SSL/TLS Decryption Controls",                  ["8.1", "8.2", "8.3"]),
        ("4.4", "Security Subscriptions",                       ["SUBS-1"]),
    ]),
    ("5", "Software Version and Lifecycle Risk", [
        ("5.1", "PAN-OS Version Lifecycle",                     ["VER-1"]),
    ]),
    ("6", "Additional Security, Architecture, and Governance Findings", [
        ("6.1",  "Password Policy",                             ["1.3.1", "1.3.2", "1.3.3",
                                                                  "1.3.4", "1.3.5", "1.3.6",
                                                                  "1.3.7", "1.3.8", "1.3.9",
                                                                  "1.3.10"]),
        ("6.2",  "Account Lockout and Session Timeout",         ["1.4.1", "1.4.2"]),
        ("6.3",  "Management Access Controls",                  ["1.2.1", "1.2.2", "1.2.3",
                                                                  "1.2.4", "1.2.5"]),
        ("6.4",  "Logging, Monitoring, and Alerting",           ["1.1.1.1", "1.1.1.2", "1.1.3"]),
        ("6.5",  "Login Banner",                                ["1.1.2"]),
        ("6.6",  "SNMP Security",                               ["1.5.1"]),
        ("6.7",  "Time Synchronisation and Update Verification", ["1.6.1", "1.6.2"]),
        ("6.8",  "User Identification Security",                ["2.1", "2.2", "2.3", "2.4",
                                                                  "2.5", "2.6", "2.7", "2.8"]),
        ("6.9",  "Dynamic Update Schedule",                     ["4.1", "4.2"]),
        ("6.10", "Certificate and VPN Security",                ["1.6.3", "1.7.1"]),
        ("6.11", "Failover Testing Plan",                       ["GOV-2"]),
        ("6.12", "Change Management and Governance",            ["GOV-4"]),
    ]),
]

_PRIORITY_COLORS = {
    "P1": COLORS["fail"],
    "P2": COLORS["critical"],
    "P3": COLORS["manual"],
    "P4": COLORS["low"],
}


def generate_consulting_docx(data: dict, output_path: str) -> str:
    """
    Consulting-style Word report following data-source/template_1.txt.
    Shows FAIL and RECOMMENDATION findings only. Empty sections are skipped.
    RECOMMENDATION findings are never included in the Section 7 priority plan.
    Vendor-dispatched: uses _PANOS_CONSULTING_MAP for palo_alto, _CONSULTING_MAP otherwise.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx required: uv add python-docx")

    p    = _prepare(data)
    meta = p["meta"]
    narr = p["narrative"]
    doc  = _DocxDocument()

    vendor         = meta.get("device_context", {}).get("vendor", "checkpoint")
    consulting_map = _PANOS_CONSULTING_MAP if vendor == "palo_alto" else _CONSULTING_MAP

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── local helpers ────────────────────────────────────────────────────────
    def h1(text):
        doc.add_heading(text, level=1)

    def h2(text):
        doc.add_heading(text, level=2)

    def h3(text):
        doc.add_heading(text, level=3)

    def body(text, bold=False, italic=False):
        para = doc.add_paragraph()
        run  = para.add_run(_strip_html(str(text)))
        run.bold   = bold
        run.italic = italic
        return para

    def bold_label(text):
        para = doc.add_paragraph()
        run  = para.add_run(str(text))
        run.bold = True
        run.font.size = Pt(10)
        return para

    def bullet(text):
        doc.add_paragraph(_strip_html(str(text)), style="List Bullet")

    def spacer():
        doc.add_paragraph()

    # ── index findings by control_id ─────────────────────────────────────────
    findings_by_id: dict = {}
    for f in p["findings"]:
        findings_by_id[f.get("control_id", "")] = f

    def section_findings(check_ids):
        """Return (fails, recs) lists for a set of check IDs."""
        fails, recs = [], []
        for cid in check_ids:
            f = findings_by_id.get(cid)
            if f is None:
                continue
            if f["status"] == "FAIL":
                fails.append(f)
            elif f["status"] == "RECOMMENDATION":
                recs.append(f)
        return fails, recs

    def risk_level(finding):
        rl = (finding.get("ai_analysis", {}).get("risk_level", "") or "").strip()
        if rl in ("Critical", "High", "Medium", "Low"):
            return rl
        # Fall back to check-supplied default when AI hasn't run
        default = (finding.get("default_risk_level", "") or "").strip()
        return default if default in ("Critical", "High", "Medium", "Low") else "Medium"

    # Accumulates (label, cis_id, description, status) for the CIS mapping appendix
    _cis_map_entries: list = []

    def render_fail(finding, label: str):
        cid = finding.get("control_id", "")
        ai  = finding.get("ai_analysis", {})
        rl  = risk_level(finding)
        actual_raw   = str(finding.get("actual", "") or "")
        actual_lines = [' '.join(l.split()) for l in actual_raw.split('\n') if l.strip()]
        actual = ' | '.join(actual_lines[:2])
        if len(actual) > 220:
            actual = actual[:220] + '...'
        check_risk = _strip_html(str(finding.get("risk_description", "") or ""))
        impact = _strip_html(str(ai.get("business_impact", "") or check_risk or ""))
        attack = _strip_html(str(ai.get("attack_scenario", "") or ""))
        rem    = _strip_html(str(ai.get("remediation_steps", "") or finding.get("remediation", "") or ""))
        if isinstance(ai.get("remediation_steps"), list):
            rem = ". ".join(ai["remediation_steps"])

        h3(f"{label} {finding.get('description', '')}")
        if actual:
            body(f"Evidence: {actual}", italic=True)
        img_b64 = finding.get("evidence_image")
        if img_b64:
            try:
                import base64
                img_bytes = base64.b64decode(img_b64)
                doc.add_picture(io.BytesIO(img_bytes), width=Cm(14))
            except Exception:
                pass
        bold_label(f"Risk Level — {rl}")
        bold_label("Risk:")
        if impact:
            bullet(impact)
        if attack:
            bullet(f"Attack scenario: {attack}")
        if not impact and not attack:
            bullet("Failure to comply with this control increases exposure to security incidents.")
        bold_label("Recommendation:")
        import re as _re2
        # Strip numbered step instructions (1. Go to... 2. Click...) — consulting reports
        # state WHAT to do, not the step-by-step HOW. Take only content before first step.
        rem_high_level = _re2.split(r'\s*\b\d+\.\s+(?=[A-Z])', rem)[0].strip()
        rem_src = rem_high_level if rem_high_level else rem
        rem_lines = [l.strip() for l in rem_src.split(". ") if l.strip() and len(l.strip()) > 3]
        if rem_lines:
            for line in rem_lines:
                bullet(line if line.endswith(".") else line + ".")
        else:
            bullet("Refer to the CIS benchmark guidance for detailed remediation steps.")
        spacer()
        _cis_map_entries.append((label, cid, finding.get("description", ""), "FAIL"))

    def render_rec(finding, label: str):
        cid   = finding.get("control_id", "")
        rem   = _strip_html(str(finding.get("remediation", "") or ""))
        notes = _strip_html(str(finding.get("notes", "") or ""))
        import re as _re
        notes = _re.sub(r':\s*\.\s+', ': ', notes).strip()
        h3(f"{label} {finding.get('description', '')}")
        if notes:
            body(notes, italic=True)
        rem_lines = [l.strip() for l in rem.split(". ") if l.strip()]
        for line in rem_lines:
            bullet(line if line.endswith(".") else line + ".")
        spacer()
        _cis_map_entries.append((label, cid, finding.get("description", ""), "RECOMMENDATION"))

    # ── Cover page ───────────────────────────────────────────────────────────
    ctx    = meta.get("device_context", {})
    target = meta.get("target", "Unknown")
    org    = ctx.get("organization", "") or target
    ts     = (meta.get("generated", "")[:10] or
              __import__("datetime").datetime.utcnow().strftime("%Y-%m-%d"))

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{org} — Firewall Security Audit")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x25, 0x2F)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.add_run("Security Assessment Report — CONFIDENTIAL").font.size = Pt(13)

    spacer()

    cover_rows = [
        ("Target Device",   target),
        ("Assessment Date", ts),
        ("Organization",    org),
        ("Framework",       meta.get("benchmark", "CIS / FW AI Audit Security Assessment")),
        ("Classification",  "CONFIDENTIAL"),
    ]
    cover_tbl = doc.add_table(rows=len(cover_rows), cols=2)
    cover_tbl.style = "Table Grid"
    for ri, (lbl, val) in enumerate(cover_rows):
        cells = cover_tbl.rows[ri].cells
        _docx_set_cell_bg(cells[0], COLORS["light"])
        _docx_cell_text(cells[0], lbl, bold=True, size_pt=10)
        _docx_cell_text(cells[1], val, size_pt=10)

    doc.add_page_break()

    # ── Table of Contents ────────────────────────────────────────────────────
    h1("Table of Contents")
    _docx_add_toc(doc)
    doc.add_page_break()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    h1("1. Executive Summary")

    exec_s = narr.get("executive_summary", "")
    if isinstance(exec_s, dict):
        headline = exec_s.get("headline", "")
        if headline:
            hl_para = doc.add_paragraph()
            hl_run  = hl_para.add_run(_strip_html(headline))
            hl_run.bold   = True
            hl_run.italic = True
        for key in ["paragraph_1", "paragraph_2", "paragraph_3"]:
            txt = exec_s.get(key, "")
            if txt:
                body(_strip_html(txt))
    elif isinstance(exec_s, str) and exec_s.strip():
        body(_strip_html(exec_s))
    else:
        body(
            f"This report presents the findings of a firewall security assessment conducted "
            f"against {target}. The assessment evaluated security configuration against "
            f"industry benchmarks and compliance best practices. This is a point-in-time "
            f"assessment and should be reviewed in the context of the organisation's broader "
            f"security programme."
        )

    spacer()
    fails_all = [f for f in p["findings"] if f["status"] == "FAIL"]
    recs_all  = [f for f in p["findings"] if f["status"] == "RECOMMENDATION"]

    if fails_all or recs_all:
        bold_label("Key findings from this assessment:")
        top_fails = sorted(
            fails_all,
            key=lambda x: int(x.get("ai_analysis", {}).get("priority_rank", 99) or 99)
        )[:6]
        for f in top_fails:
            bullet(f"✗  {f['description']}")
        for r in recs_all[:3]:
            bullet(f"▶  {r['description']}")

    doc.add_page_break()

    # ── Sections 2–6 from consulting_map (vendor-dispatched) ─────────────────
    for sec_num, sec_title, subsections in consulting_map:

        # Collect subsections that have content
        populated = []
        for sub_num, sub_title, check_ids in subsections:
            fails, recs = section_findings(check_ids)
            if fails or recs:
                populated.append((sub_num, sub_title, fails, recs))

        if not populated:
            continue

        h1(f"{sec_num}. {sec_title}")

        for sub_num, sub_title, fails, recs in populated:
            h2(f"{sub_num} {sub_title}")
            sub_n = [0]
            for finding in fails:
                sub_n[0] += 1
                render_fail(finding, f"{sub_num}.{sub_n[0]}")
            for rec in recs:
                sub_n[0] += 1
                render_rec(rec, f"{sub_num}.{sub_n[0]}")

        doc.add_page_break()

    # ── 7. Initiative List & Priority Plan ───────────────────────────────────
    h1("7. Initiative List and Priority Plan")

    # ── 7.1 Prioritisation Framework ─────────────────────────────────────────
    h2("7.1 Prioritisation Framework")
    body(
        "To organise and implement firewall initiatives, an Impact vs. Effort "
        "prioritisation matrix was used. Each initiative is evaluated against two axes:"
    )
    spacer()

    bold_label("Impact")
    bullet(
        "High: Significantly reduces enterprise risk, ensures high availability and "
        "resilience, addresses critical audit or compliance gaps, or prevents material "
        "security incidents."
    )
    bullet(
        "Medium: Strengthens the security posture, improves operational efficiency and "
        "visibility, and aligns configurations and practices with good security standards."
    )
    bullet(
        "Low: Delivers incremental hygiene, optimisation, or documentation improvements "
        "with limited standalone effect on overall risk."
    )
    spacer()

    bold_label("Effort")
    bullet(
        "High: Requires substantial resources, complex configuration changes, or downtime; "
        "may involve multiple teams or extended planning."
    )
    bullet(
        "Medium: Moderate resource allocation and planning; manageable within standard "
        "maintenance windows with minimal disruption."
    )
    bullet(
        "Low: Simple changes with minimal resource use and no significant operational impact; "
        "can typically be completed within a single maintenance window."
    )
    spacer()

    body(
        "Initiatives with High Impact and Low Effort are considered quick wins and "
        "prioritised first, while those with Low Impact and High Effort are addressed last."
    )
    spacer()

    matrix = doc.add_table(rows=4, cols=4)
    matrix.style = "Table Grid"
    matrix_data = [
        ("Impact \\ Effort", "High Effort",                  "Medium Effort",             "Low Effort"),
        ("High Impact",      "P3 — Strategic",               "P2 — Important",            "P1 — Quick Wins"),
        ("Medium Impact",    "P4 — Ad-hoc Improvements",     "P3 — Strategic",            "P2 — Important"),
        ("Low Impact",       "P4 — Ad-hoc Improvements",     "P4 — Ad-hoc Improvements", "P3 — Strategic"),
    ]
    for ri, row_data in enumerate(matrix_data):
        for ci, val in enumerate(row_data):
            cell = matrix.rows[ri].cells[ci]
            is_header = ri == 0 or ci == 0
            _docx_set_cell_bg(cell, COLORS["header"] if is_header else COLORS["white"])
            _docx_cell_text(cell, val, bold=is_header,
                            color_hex=COLORS["white"] if is_header else COLORS["header"],
                            size_pt=9)
    spacer()

    # ── 7.2 Initiatives ───────────────────────────────────────────────────────
    h2("7.2 Initiatives")
    body(
        "Building on the prioritisation framework defined in section 7.1, this section "
        "consolidates the audit recommendations into structured initiatives, each "
        "representing a coherent set of related findings that can be addressed together."
    )
    body(
        "The initiatives identified balance risk reduction with delivery practicality, "
        "enabling a sequenced plan for delivering near-term security improvements and "
        "longer-term structural changes."
    )
    spacer()

    # ── 7.2.1 Initiatives Quadrant ────────────────────────────────────────────
    h2("7.2.1 Initiatives Quadrant")
    body(
        "The initiatives are positioned on an impact vs. effort matrix to reflect their "
        "expected contribution to risk reduction and the resources required for implementation."
    )
    spacer()

    quad_matrix = doc.add_table(rows=4, cols=4)
    quad_matrix.style = "Table Grid"
    for ri, row_data in enumerate(matrix_data):
        for ci, val in enumerate(row_data):
            cell = quad_matrix.rows[ri].cells[ci]
            is_header = ri == 0 or ci == 0
            _docx_set_cell_bg(cell, COLORS["header"] if is_header else COLORS["light"])
            _docx_cell_text(cell, val, bold=is_header,
                            color_hex=COLORS["white"] if is_header else COLORS["header"],
                            size_pt=9)
    spacer()

    # ── 7.2.2 Initiatives Summary Table ──────────────────────────────────────
    h2("7.2.2 Initiatives Summary")
    body(
        "The table below provides a high-level index of the initiatives, including the "
        "initiative ID, name, source finding chapter, CIS Benchmark reference, and "
        "impact/effort/priority classification."
    )
    spacer()

    def _find_source_chapter(cid):
        """Return just the subsection number (e.g. '3.1') for the Finding Chapter column."""
        for sn, _, subs in consulting_map:
            for sub_n, sub_t, ids in subs:
                if cid in ids:
                    return sub_n
        return "—"

    def _derive_priority(finding):
        ai         = finding.get("ai_analysis", {})
        rl         = (ai.get("risk_level", "") or "").strip()
        effort_raw = (ai.get("remediation_effort", "") or "").lower()
        if "low"  in effort_raw:
            effort = "Low"
        elif "high" in effort_raw:
            effort = "High"
        else:
            effort = "Medium"
        if rl in ("Critical", "High"):
            impact   = "High"
            priority = "P1" if effort == "Low" else ("P2" if effort == "Medium" else "P3")
        elif rl == "Medium":
            impact   = "Medium"
            priority = "P2" if effort == "Low" else ("P3" if effort == "Medium" else "P4")
        else:
            impact   = "Low"
            priority = "P3" if effort in ("Low", "Medium") else "P4"
        return impact, effort, priority

    if not fails_all:
        body("No failed findings were identified in this assessment.")
    else:
        sorted_fails = sorted(
            fails_all,
            key=lambda x: int(x.get("ai_analysis", {}).get("priority_rank", 99) or 99)
        )
        headers = ["#", "Initiative", "Finding Chapter", "CIS Benchmark", "Impact", "Effort", "Priority"]
        col_w   = [Cm(0.8), Cm(6.5), Cm(2.2), Cm(2.2), Cm(1.5), Cm(1.5), Cm(1.5)]
        tbl     = doc.add_table(rows=1 + len(sorted_fails), cols=len(headers))
        tbl.style = "Table Grid"

        hdr_row = tbl.rows[0]
        for ci, (h, w) in enumerate(zip(headers, col_w)):
            hdr_row.cells[ci].width = w
            _docx_set_cell_bg(hdr_row.cells[ci], COLORS["header"])
            _docx_cell_text(hdr_row.cells[ci], h, bold=True,
                            color_hex=COLORS["white"], size_pt=9,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        for ri, finding in enumerate(sorted_fails, 1):
            impact, effort, priority = _derive_priority(finding)
            cid     = finding.get("control_id", "—")
            chapter = _find_source_chapter(cid)
            vals    = [str(ri), finding.get("description", ""),
                       chapter, cid, impact, effort, priority]
            for ci, (val, w) in enumerate(zip(vals, col_w)):
                cell = tbl.rows[ri].cells[ci]
                cell.width = w
                if ci == 6:
                    _docx_set_cell_bg(cell, _PRIORITY_COLORS.get(priority, COLORS["mid"]))
                    _docx_cell_text(cell, val, bold=True,
                                    color_hex=COLORS["white"], size_pt=9,
                                    align=WD_ALIGN_PARAGRAPH.CENTER)
                elif ci in (0, 2, 3, 4, 5):
                    bg = COLORS["light"] if ri % 2 == 0 else COLORS["white"]
                    _docx_set_cell_bg(cell, bg)
                    _docx_cell_text(cell, val, size_pt=9,
                                    align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    bg = COLORS["light"] if ri % 2 == 0 else COLORS["white"]
                    _docx_set_cell_bg(cell, bg)
                    _docx_cell_text(cell, val, size_pt=9)

    # ── Appendix: CIS Benchmark Reference ────────────────────────────────────
    if _cis_map_entries:
        doc.add_page_break()
        h1("Appendix: CIS Benchmark Reference")
        body(
            "The table below maps each finding in this report to its corresponding "
            "CIS Benchmark check ID, enabling direct cross-reference with the benchmark document."
        )
        spacer()

        headers  = ["Finding Ref", "CIS Check ID", "Description", "Type"]
        col_w    = [Cm(2.2), Cm(2.4), Cm(10.5), Cm(2.5)]
        apdx_tbl = doc.add_table(rows=1 + len(_cis_map_entries), cols=4)
        apdx_tbl.style = "Table Grid"

        hdr_row = apdx_tbl.rows[0]
        for ci, (h, w) in enumerate(zip(headers, col_w)):
            hdr_row.cells[ci].width = w
            _docx_set_cell_bg(hdr_row.cells[ci], COLORS["header"])
            _docx_cell_text(hdr_row.cells[ci], h, bold=True,
                            color_hex=COLORS["white"], size_pt=9,
                            align=WD_ALIGN_PARAGRAPH.CENTER)

        for ri, (label, cid, desc, status) in enumerate(_cis_map_entries, 1):
            bg = COLORS["light"] if ri % 2 == 0 else COLORS["white"]
            row = apdx_tbl.rows[ri]
            for ci, (val, w) in enumerate(zip([label, cid, desc, status], col_w)):
                cell = row.cells[ci]
                cell.width = w
                _docx_set_cell_bg(cell, bg)
                align = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
                color = (COLORS["fail"] if status == "FAIL" else
                         COLORS["info"] if status == "RECOMMENDATION" else COLORS["header"])
                _docx_cell_text(cell, val, size_pt=9, align=align,
                                bold=(ci == 0),
                                color_hex=color if ci == 3 else COLORS["header"])

    doc.save(output_path)
    print(f"[REPORT] ✓ CONSULTING → {os.path.abspath(output_path)}")
    return output_path


def generate_all(data: dict, base_path: str) -> dict:
    """Generate CSV, Excel, and PDF from a single enriched data dict."""
    paths = {}
    csv_path   = base_path + ".csv"
    excel_path = base_path + ".xlsx"
    pdf_path   = base_path + ".pdf"

    paths["csv"]   = generate_csv(data, csv_path)
    if HAS_OPENPYXL:
        paths["excel"] = generate_excel(data, excel_path)
    else:
        print("[REPORT] ⚠️  openpyxl not available — Excel skipped.")
    if HAS_REPORTLAB:
        paths["pdf"]   = generate_pdf(data, pdf_path)
    else:
        print("[REPORT] ⚠️  reportlab not available — PDF skipped.")
    return paths


# ============================================================================
# CLI
# ============================================================================
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate CSV/Excel/PDF reports from audit JSON")
    parser.add_argument("audit_json", help="Enriched or raw audit JSON file")
    parser.add_argument("--output-base", "-o", default=None,
                        help="Base path for output files (no extension)")
    parser.add_argument("--format", choices=["csv", "excel", "pdf", "all"],
                        default="all", help="Output format (default: all)")
    args = parser.parse_args()

    with open(args.audit_json) as f:
        audit_data = json.load(f)

    base = args.output_base or os.path.splitext(args.audit_json)[0] + "_report"

    if args.format == "csv":
        generate_csv(audit_data, base + ".csv")
    elif args.format == "excel":
        generate_excel(audit_data, base + ".xlsx")
    elif args.format == "pdf":
        generate_pdf(audit_data, base + ".pdf")
    else:
        generate_all(audit_data, base)
